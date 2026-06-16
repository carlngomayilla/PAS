from __future__ import annotations

import math
import shutil
import textwrap
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "manuel-utilisation-e-pilotage-pas-assets"
DOCX_PATH = DOCS / "MANUEL_UTILISATION_E_PILOTAGE_PAS.docx"
MD_PATH = DOCS / "MANUEL_UTILISATION_E_PILOTAGE_PAS.md"
LOGO_PATH = ROOT / "public" / "images" / "logo-full.png"

TITLE = "Manuel d'utilisation de l'application e-Pilotage PAS"
SUBTITLE = "Guide détaillé et illustré pour la planification, le suivi, la validation et le reporting"
VERSION = "Version 1.0 - 14 juin 2026"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
MUTED = RGBColor(85, 85, 85)
BLACK = RGBColor(0, 0, 0)
WHITE = RGBColor(255, 255, 255)
LIGHT_BLUE_HEX = "E8EEF5"
LIGHT_GRAY_HEX = "F2F4F7"
CALLOUT_HEX = "F4F6F9"
WARNING_HEX = "FFF4D6"


def font_path(name: str) -> str | None:
    for folder in [Path("C:/Windows/Fonts"), Path("/usr/share/fonts/truetype/dejavu")]:
        candidate = folder / name
        if candidate.exists():
            return str(candidate)
    return None


def get_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = []
    if bold:
        candidates.extend(["arialbd.ttf", "DejaVuSans-Bold.ttf"])
    candidates.extend(["arial.ttf", "DejaVuSans.ttf"])
    for candidate in candidates:
        path = font_path(candidate)
        if path:
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def text_size(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont) -> tuple[int, int]:
    bbox = draw.textbbox((0, 0), text, font=font)
    return bbox[2] - bbox[0], bbox[3] - bbox[1]


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    lines: list[str] = []
    for paragraph in text.split("\n"):
        words = paragraph.split()
        if not words:
            lines.append("")
            continue
        current: list[str] = []
        for word in words:
            trial = " ".join(current + [word])
            width, _ = text_size(draw, trial, font)
            if width <= max_width or not current:
                current.append(word)
            else:
                lines.append(" ".join(current))
                current = [word]
        if current:
            lines.append(" ".join(current))
    return lines


def centered_text(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.ImageFont,
    fill: str,
    line_gap: int = 6,
) -> None:
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font, x2 - x1 - 28)
    heights = [text_size(draw, line, font)[1] for line in lines]
    total_height = sum(heights) + line_gap * max(0, len(lines) - 1)
    y = y1 + (y2 - y1 - total_height) / 2
    for line, height in zip(lines, heights):
        width, _ = text_size(draw, line, font)
        draw.text((x1 + (x2 - x1 - width) / 2, y), line, font=font, fill=fill)
        y += height + line_gap


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str, width: int = 5) -> None:
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 16
    points = [
        end,
        (end[0] - size * math.cos(angle - math.pi / 6), end[1] - size * math.sin(angle - math.pi / 6)),
        (end[0] - size * math.cos(angle + math.pi / 6), end[1] - size * math.sin(angle + math.pi / 6)),
    ]
    draw.polygon(points, fill=color)


def rounded_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    fill: str,
    outline: str = "#C8D7E6",
    radius: int = 24,
    width: int = 3,
) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def save_flow_diagram(path: Path) -> None:
    image = Image.new("RGB", (1800, 560), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    title_font = get_font(38, True)
    node_font = get_font(24, True)
    small_font = get_font(18)
    draw.text((60, 38), "Chaîne complète de pilotage", font=title_font, fill="#0B2545")
    draw.text(
        (60, 88),
        "Chaque niveau prépare le suivant : la stratégie descend vers l'exécution, puis les données remontent vers le reporting.",
        font=small_font,
        fill="#555555",
    )
    nodes = [
        ("PAS\nVision stratégique", "#E8F3FB"),
        ("PAO\nDirection", "#EAF6E1"),
        ("PTA\nService", "#FFF4D6"),
        ("Actions\nExécution", "#F7E9FF"),
        ("Sous-actions\nRMO", "#EDF1FF"),
        ("Suivi\nPreuves", "#E6F8F1"),
        ("Validation\nChef", "#FFECE5"),
        ("KPI\nOfficiel", "#F0F6FF"),
        ("Reporting\nExports", "#EEF2F6"),
        ("Audit\nTraçabilité", "#F7F7F7"),
    ]
    box_w, box_h = 290, 105
    top_row = [(60 + col * 340, 170, 60 + col * 340 + box_w, 170 + box_h) for col in range(5)]
    bottom_row_left_to_right = [(60 + col * 340, 340, 60 + col * 340 + box_w, 340 + box_h) for col in range(5)]
    # Snake reading: PAS -> ... -> Sous-actions, then down to Suivi and back left-to-right in the
    # logical sequence visible to the reader.
    positions = top_row + list(reversed(bottom_row_left_to_right))
    for idx, (label, fill) in enumerate(nodes):
        box = positions[idx]
        rounded_box(draw, box, fill=fill)
        centered_text(draw, box, label, node_font, "#17324A")
        if idx < len(nodes) - 1:
            current = box
            nxt = positions[idx + 1]
            if idx == 4:
                draw_arrow(draw, (current[0] + box_w // 2, current[3] + 12), (nxt[0] + box_w // 2, nxt[1] - 12), "#3996D3")
            elif current[0] > nxt[0] and current[1] == nxt[1]:
                draw_arrow(draw, (current[0] - 10, current[1] + box_h // 2), (nxt[2] + 10, nxt[1] + box_h // 2), "#3996D3")
            else:
                draw_arrow(draw, (current[2] + 10, current[1] + box_h // 2), (nxt[0] - 10, nxt[1] + box_h // 2), "#3996D3")
    image.save(path)


def save_navigation_diagram(path: Path) -> None:
    image = Image.new("RGB", (1600, 900), "#F8FBFE")
    draw = ImageDraw.Draw(image)
    title_font = get_font(38, True)
    section_font = get_font(25, True)
    item_font = get_font(22)
    draw.text((60, 42), "Carte de navigation de l'espace de travail", font=title_font, fill="#0B2545")
    sections = [
        ("Menu", ["Pilotage", "Mes tâches", "Notifications"]),
        ("Planification", ["PAS", "PAO", "PTA", "Imports Excel"]),
        ("Exécution", ["Actions", "Financement des actions"]),
        ("Pilotage", ["Reporting", "Alertes"]),
        ("Administration", ["Référentiels", "Délégations", "Rétention", "API Docs", "Audit"]),
        ("Plateforme", ["Super Administration"]),
    ]
    colors = ["#E8F3FB", "#FFF4D6", "#EAF6E1", "#EDF1FF", "#F4F6F9", "#F7E9FF"]
    card_w, card_h = 460, 310
    for idx, (section, items) in enumerate(sections):
        row, col = divmod(idx, 3)
        x = 60 + col * 510
        y = 130 + row * 360
        rounded_box(draw, (x, y, x + card_w, y + card_h), colors[idx], "#C9D8E6")
        draw.text((x + 26, y + 24), section, font=section_font, fill="#17324A")
        yy = y + 72
        for item in items:
            draw.rounded_rectangle((x + 26, yy, x + card_w - 26, yy + 38), radius=18, fill="#FFFFFF", outline="#D8ECF8")
            draw.text((x + 46, yy + 7), item, font=item_font, fill="#1C203D")
            yy += 44
    image.save(path)


def save_roles_diagram(path: Path) -> None:
    image = Image.new("RGB", (1700, 900), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    title_font = get_font(36, True)
    header_font = get_font(21, True)
    cell_font = get_font(18)
    draw.text((60, 38), "Périmètres et responsabilités par rôle", font=title_font, fill="#0B2545")
    roles = [
        ("Super Admin", "Global", "Paramétrage complet"),
        ("Admin / Admin fonctionnel", "Global", "Organisation, référentiels, gouvernance"),
        ("DG", "Global", "Supervision, arbitrages, lecture consolidée"),
        ("Planification / SCIQ", "Global", "Structuration, contrôle transversal, reporting"),
        ("Direction", "Direction", "PAO, supervision des PTA et actions"),
        ("Service / Chef unité", "Service", "PTA, création d'actions, validation chef"),
        ("Agent / RMO", "Assigné", "Suivi, preuves, commentaires, soumission"),
        ("Cabinet / Collaborateur", "Global lecture", "Consultation et suivi institutionnel"),
    ]
    x0, y0 = 60, 120
    widths = [380, 270, 830]
    row_h = 76
    headers = ["Rôle", "Périmètre", "Contribution principale"]
    x = x0
    for w, header in zip(widths, headers):
        draw.rectangle((x, y0, x + w, y0 + row_h), fill="#17324A")
        centered_text(draw, (x, y0, x + w, y0 + row_h), header, header_font, "#FFFFFF")
        x += w
    for idx, row in enumerate(roles, start=1):
        y = y0 + idx * row_h
        fill = "#F8FBFE" if idx % 2 else "#FFFFFF"
        x = x0
        for col, (w, value) in enumerate(zip(widths, row)):
            draw.rectangle((x, y, x + w, y + row_h), fill=fill, outline="#D8ECF8")
            font = header_font if col == 0 else cell_font
            centered_text(draw, (x + 8, y, x + w - 8, y + row_h), value, font, "#1C203D")
            x += w
    image.save(path)


def save_action_workflow(path: Path) -> None:
    image = Image.new("RGB", (1700, 980), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    title_font = get_font(36, True)
    box_font = get_font(21, True)
    note_font = get_font(17)
    draw.text((60, 38), "Workflow d'une action : provisoire puis officielle", font=title_font, fill="#0B2545")
    lanes = [
        ("PTA / Chef", 70, "#FFF4D6"),
        ("Agent / RMO", 470, "#E8F3FB"),
        ("Chef validateur", 870, "#EAF6E1"),
        ("Reporting", 1270, "#F4F6F9"),
    ]
    for label, x, fill in lanes:
        draw.rounded_rectangle((x, 105, x + 340, 905), radius=20, fill=fill, outline="#D8ECF8", width=3)
        centered_text(draw, (x + 20, 120, x + 320, 168), label, box_font, "#17324A")
    steps = [
        ((100, 220, 380, 310), "Définir l'action\nType, cible, seuils,\nresponsable, preuves"),
        ((500, 220, 780, 310), "Enregistrer\nBrouillon libre,\nperformance provisoire"),
        ((500, 410, 780, 500), "Soumettre\nContrôles complets :\npreuve, commentaire, difficulté"),
        ((900, 410, 1180, 500), "Valider ou rejeter\nMotif obligatoire\nen cas de rejet"),
        ((900, 610, 1180, 700), "Officialiser\nProgression figée\npar le chef"),
        ((1300, 610, 1580, 700), "Consolider\nKPI, tableaux,\nexports PDF/Excel"),
    ]
    for box, label in steps:
        rounded_box(draw, box, "#FFFFFF", "#9CCCE8")
        centered_text(draw, box, label, note_font, "#1C203D")
    draw_arrow(draw, (380, 265), (500, 265), "#3996D3")
    draw_arrow(draw, (640, 310), (640, 410), "#3996D3")
    draw_arrow(draw, (780, 455), (900, 455), "#3996D3")
    draw_arrow(draw, (1040, 500), (1040, 610), "#8FC043")
    draw_arrow(draw, (1180, 655), (1300, 655), "#3996D3")
    draw_arrow(draw, (900, 500), (780, 500), "#F9B13C")
    draw.text((650, 520), "Rejet : retour à correction", font=note_font, fill="#7A5A00")
    draw.text((70, 930), "Règle clé : les chiffres du reporting s'appuient sur la performance officielle, pas sur le brouillon de suivi.", font=get_font(20, True), fill="#17324A")
    image.save(path)


def save_action_creation(path: Path) -> None:
    image = Image.new("RGB", (1500, 980), "#F8FBFE")
    draw = ImageDraw.Draw(image)
    title_font = get_font(36, True)
    step_font = get_font(22, True)
    body_font = get_font(18)
    draw.text((60, 42), "Création d'une action dans le PTA", font=title_font, fill="#0B2545")
    steps = [
        ("1. Identification", "Libellé de l'action et informations générales."),
        ("2. Responsable / affectation", "Choisir un ou plusieurs RMO / agents assignés."),
        ("3. Planification", "Date de début et date de fin, dans l'échéance de l'objectif opérationnel."),
        ("4. Cible et seuil", "Type d'action, cible, unité, seuil unique ou trimestriel."),
        ("5. Sous-actions prévues", "Uniquement pour une action composée ; les poids doivent totaliser 100 %."),
        ("6. Ressources nécessaires", "Main-d'œuvre, matériel, ressources techniques ou détails complémentaires."),
        ("7. Risques", "Risque potentiel, niveau de risque et mesures préventives."),
        ("8. Financement", "Besoin de financement, montant, nature et pièce justificative si nécessaire."),
    ]
    x, y = 85, 135
    for idx, (heading, detail) in enumerate(steps):
        fill = "#FFFFFF" if idx % 2 == 0 else "#E8F3FB"
        rounded_box(draw, (x, y, x + 1330, y + 82), fill, "#C8D7E6", radius=18)
        draw.text((x + 26, y + 18), heading, font=step_font, fill="#17324A")
        draw.text((x + 360, y + 20), detail, font=body_font, fill="#555555")
        if idx < len(steps) - 1:
            draw_arrow(draw, (x + 665, y + 86), (x + 665, y + 112), "#3996D3", width=4)
        y += 102
    image.save(path)


def save_reporting_alerts(path: Path) -> None:
    image = Image.new("RGB", (1600, 820), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    title_font = get_font(36, True)
    box_font = get_font(23, True)
    body_font = get_font(18)
    draw.text((60, 40), "Reporting, alertes et exports", font=title_font, fill="#0B2545")
    rounded_box(draw, (610, 180, 990, 360), "#E8F3FB", "#9CCCE8")
    centered_text(draw, (630, 200, 970, 340), "Base officielle\nActions validées\net périmètre autorisé", box_font, "#17324A")
    boxes = [
        ((90, 180, 430, 330), "Filtres", "Exercice, trimestre, direction,\nservice, statut, type d'action."),
        ((90, 480, 430, 630), "Alertes", "Retards, criticités, anomalies,\nnotifications et lecture directe."),
        ((1170, 180, 1510, 330), "Export Excel", "Données consolidées, tables,\nfiltrage et contrôle."),
        ((1170, 480, 1510, 630), "Export PDF", "Rapport imprimable pour\ndiffusion institutionnelle."),
    ]
    for box, heading, detail in boxes:
        rounded_box(draw, box, "#FFFFFF", "#C8D7E6")
        centered_text(draw, (box[0] + 12, box[1] + 18, box[2] - 12, box[1] + 58), heading, box_font, "#17324A")
        centered_text(draw, (box[0] + 20, box[1] + 68, box[2] - 20, box[3] - 20), detail, body_font, "#555555")
    draw_arrow(draw, (430, 255), (610, 255), "#3996D3")
    draw_arrow(draw, (990, 255), (1170, 255), "#3996D3")
    draw_arrow(draw, (430, 555), (610, 320), "#F9B13C")
    draw_arrow(draw, (990, 320), (1170, 555), "#3996D3")
    draw.text((60, 720), "À retenir : un indicateur de suivi n'est consolidé que si le circuit de validation requis est terminé.", font=get_font(22, True), fill="#17324A")
    image.save(path)


def save_super_admin(path: Path) -> None:
    image = Image.new("RGB", (1600, 900), "#F8FBFE")
    draw = ImageDraw.Draw(image)
    title_font = get_font(36, True)
    group_font = get_font(26, True)
    item_font = get_font(19)
    draw.text((60, 42), "Super Administration : zones de paramétrage", font=title_font, fill="#0B2545")
    groups = [
        ("Plateforme", ["Généraux", "Apparence", "Modules", "Maintenance"], "#E8F3FB"),
        ("Gouvernance", ["Rôles", "Organisation", "Unités DG", "Dashboards", "Diagnostic", "Audit"], "#EAF6E1"),
        ("Pilotage", ["Workflow", "Exercices", "Calcul", "Actions", "Référentiels", "Documents", "KPI", "Notifications"], "#FFF4D6"),
        ("Avancé", ["Snapshots", "Simulation", "Templates d'export"], "#F7E9FF"),
    ]
    positions = [(70, 135), (830, 135), (70, 500), (830, 500)]
    for (heading, items, fill), (x, y) in zip(groups, positions):
        rounded_box(draw, (x, y, x + 680, y + 305), fill, "#C8D7E6")
        draw.text((x + 28, y + 24), heading, font=group_font, fill="#17324A")
        yy = y + 82
        for item in items:
            draw.rounded_rectangle((x + 28, yy, x + 318, yy + 38), radius=16, fill="#FFFFFF", outline="#D8ECF8")
            draw.text((x + 44, yy + 8), item, font=item_font, fill="#1C203D")
            yy += 44
            if yy > y + 252:
                yy = y + 82
                x += 320
    image.save(path)


def create_assets() -> dict[str, Path]:
    ASSETS.mkdir(parents=True, exist_ok=True)
    assets = {
        "chaine_pilotage": ASSETS / "schema-chaine-pilotage.png",
        "navigation": ASSETS / "schema-navigation-modules.png",
        "roles": ASSETS / "schema-roles-perimetres.png",
        "workflow_action": ASSETS / "schema-workflow-action.png",
        "creation_action": ASSETS / "schema-creation-action.png",
        "reporting_alertes": ASSETS / "schema-reporting-alertes.png",
        "super_admin": ASSETS / "schema-super-admin.png",
    }
    save_flow_diagram(assets["chaine_pilotage"])
    save_navigation_diagram(assets["navigation"])
    save_roles_diagram(assets["roles"])
    save_action_workflow(assets["workflow_action"])
    save_action_creation(assets["creation_action"])
    save_reporting_alerts(assets["reporting_alertes"])
    save_super_admin(assets["super_admin"])
    if LOGO_PATH.exists():
        shutil.copy2(LOGO_PATH, ASSETS / "logo-e-pilotage-pas.png")
        assets["logo"] = ASSETS / "logo-e-pilotage-pas.png"
    return assets


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Pt(widths[idx] / 20)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, name: str = "Calibri", size: int | None = None, color: RGBColor | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


class ManualWriter:
    def __init__(self, doc: Document, assets: dict[str, Path]):
        self.doc = doc
        self.assets = assets
        self.md: list[str] = []

    def md_image_path(self, path: Path) -> str:
        return path.relative_to(DOCS).as_posix()

    def h1(self, text: str) -> None:
        self.doc.add_heading(text, level=1)
        self.md.extend([f"## {text}", ""])

    def h2(self, text: str) -> None:
        self.doc.add_heading(text, level=2)
        self.md.extend([f"### {text}", ""])

    def h3(self, text: str) -> None:
        self.doc.add_heading(text, level=3)
        self.md.extend([f"#### {text}", ""])

    def p(self, text: str) -> None:
        para = self.doc.add_paragraph(text)
        para.style = self.doc.styles["Normal"]
        self.md.extend([text, ""])

    def bullets(self, items: list[str]) -> None:
        for item in items:
            para = self.doc.add_paragraph(item, style="List Bullet")
            para.paragraph_format.space_after = Pt(4)
            self.md.append(f"- {item}")
        self.md.append("")

    def steps(self, items: list[str]) -> None:
        for idx, item in enumerate(items, start=1):
            para = self.doc.add_paragraph(item, style="List Number")
            para.paragraph_format.space_after = Pt(4)
            self.md.append(f"{idx}. {item}")
        self.md.append("")

    def callout(self, title: str, text: str, fill: str = CALLOUT_HEX) -> None:
        table = self.doc.add_table(rows=1, cols=1)
        set_table_width(table, [9360])
        cell = table.cell(0, 0)
        set_cell_shading(cell, fill)
        set_cell_margins(cell, 120, 160, 120, 160)
        p = cell.paragraphs[0]
        r = p.add_run(title + " : ")
        set_run_font(r, size=11, color=NAVY, bold=True)
        r2 = p.add_run(text)
        set_run_font(r2, size=11, color=BLACK)
        self.doc.add_paragraph()
        self.md.extend([f"> **{title} :** {text}", ""])

    def table(self, headers: list[str], rows: list[list[str]], widths: list[int] | None = None) -> None:
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        if widths is None:
            widths = [9360 // len(headers)] * len(headers)
            widths[-1] += 9360 - sum(widths)
        hdr_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            set_cell_shading(hdr_cells[idx], LIGHT_BLUE_HEX)
            set_cell_margins(hdr_cells[idx])
            hdr_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = hdr_cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header)
            set_run_font(run, size=10, color=NAVY, bold=True)
        for row in rows:
            cells = table.add_row().cells
            for idx, value in enumerate(row):
                set_cell_margins(cells[idx])
                cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p = cells[idx].paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(value)
                set_run_font(run, size=9, color=BLACK)
        set_table_width(table, widths)
        self.doc.add_paragraph()
        self.md.append("| " + " | ".join(headers) + " |")
        self.md.append("| " + " | ".join(["---"] * len(headers)) + " |")
        for row in rows:
            self.md.append("| " + " | ".join(value.replace("\n", "<br>") for value in row) + " |")
        self.md.append("")

    def image(self, key: str, caption: str, width: float = 6.3) -> None:
        path = self.assets[key]
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(str(path), width=Inches(width))
        cap = self.doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.style = self.doc.styles["Caption"]
        self.md.extend([f"![{caption}]({self.md_image_path(path)})", f"*{caption}*", ""])

    def page_break(self) -> None:
        self.doc.add_page_break()
        self.md.extend(["---", ""])


def configure_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = MUTED
    caption.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0]
    header.text = "e-Pilotage PAS - Manuel d'utilisation"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "ANBG - Usage interne"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run_font(run, size=9, color=MUTED)

    return doc


def add_cover(writer: ManualWriter) -> None:
    doc = writer.doc
    if "logo" in writer.assets:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(writer.assets["logo"]), width=Inches(2.3))
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(TITLE)
    set_run_font(r, size=26, color=NAVY, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run(SUBTITLE)
    set_run_font(r, size=13, color=MUTED)
    doc.add_paragraph()
    meta = doc.add_table(rows=4, cols=2)
    set_table_width(meta, [2600, 5600])
    rows = [
        ("Application", "ANBG e-Pilotage PAS / PAO / PTA / Actions"),
        ("Public", "DG, Planification, SCIQ, Directions, Services, Agents, administrateurs"),
        ("Version du manuel", VERSION),
        ("Format", "Guide utilisateur détaillé avec illustrations et procédures pas à pas"),
    ]
    for idx, (label, value) in enumerate(rows):
        cells = meta.rows[idx].cells
        set_cell_shading(cells[0], LIGHT_BLUE_HEX)
        set_cell_margins(cells[0])
        set_cell_margins(cells[1])
        run = cells[0].paragraphs[0].add_run(label)
        set_run_font(run, size=10, color=NAVY, bold=True)
        run = cells[1].paragraphs[0].add_run(value)
        set_run_font(run, size=10, color=BLACK)
    doc.add_paragraph()
    writer.callout(
        "Objectif",
        "Ce manuel explique comment utiliser l'application au quotidien : créer la planification, suivre les actions, valider les réalisations, consulter les alertes, exporter les rapports et administrer la plateforme.",
        fill=WARNING_HEX,
    )
    writer.md = [
        f"# {TITLE}",
        "",
        f"**{SUBTITLE}**",
        "",
        f"**{VERSION}**",
        "",
        "![Logo e-Pilotage PAS](manuel-utilisation-e-pilotage-pas-assets/logo-e-pilotage-pas.png)",
        "",
    ]
    writer.page_break()


def build_manual(writer: ManualWriter) -> None:
    writer.h1("1. Présentation générale")
    writer.p(
        "e-Pilotage PAS est l'application de pilotage stratégique et opérationnel de l'ANBG. Elle permet de partir du Plan d'Actions Stratégique (PAS), de le décliner en Plans d'Actions Opérationnels (PAO), puis en Plans de Travail Annuels (PTA), avant de suivre les actions, leurs justificatifs, leurs validations et leurs indicateurs."
    )
    writer.image("chaine_pilotage", "Illustration 1 - Chaîne PAS, PAO, PTA, actions, suivi, KPI et reporting")
    writer.h2("Ce que l'application permet de faire")
    writer.bullets([
        "Structurer un PAS avec ses axes stratégiques et ses objectifs stratégiques.",
        "Décliner le PAS en PAO annuels par direction et par objectif stratégique.",
        "Construire les PTA des services à partir des objectifs opérationnels transmis.",
        "Créer les actions dans le PTA, affecter les RMO ou agents, définir les cibles, les seuils, les justificatifs, les risques et les financements.",
        "Suivre les actions avec quantités réalisées, justificatifs, commentaires et difficultés.",
        "Valider ou rejeter les réalisations pour officialiser la performance.",
        "Consulter les tableaux de bord, les alertes, les rapports et les exports PDF/Excel.",
        "Administrer les rôles, les modules, les référentiels, les exercices, les règles de calcul et les modèles d'export.",
    ])
    writer.callout(
        "Principe métier",
        "Le PTA définit les règles de l'action. Le suivi applique ces règles. La validation du chef officialise la performance. Le reporting utilise la performance officielle.",
    )

    writer.h1("2. Concepts clés à connaître")
    writer.table(
        ["Concept", "Définition simple", "Exemple d'utilisation"],
        [
            ["PAS", "Plan d'Actions Stratégique pluriannuel. Il contient les axes et objectifs stratégiques.", "Définir les priorités institutionnelles sur une période donnée."],
            ["PAO", "Plan d'Actions Opérationnel annuel d'une direction.", "Décliner un objectif stratégique en objectifs opérationnels pour la DSIC, la DAF ou une autre direction."],
            ["PTA", "Plan de Travail Annuel d'un service ou d'une unité.", "Planifier les actions concrètes du service pour atteindre les objectifs reçus."],
            ["Action", "Unité opérationnelle suivie dans l'application.", "Organiser une mission, produire un rapport, réaliser une activité, livrer un indicateur."],
            ["Sous-action", "Découpage d'une action composée entre plusieurs RMO.", "Répartir une action en tâches pondérées par responsable."],
            ["Justificatif", "Pièce qui prouve l'exécution ou le financement.", "PDF, image, document Word, Excel, photo ou autre preuve autorisée."],
            ["Performance provisoire", "Progression calculée lors du suivi, avant validation.", "Un agent enregistre 80 %, mais le chef n'a pas encore validé."],
            ["Performance officielle", "Progression figée après validation du chef.", "Seule cette valeur entre dans les statistiques et le reporting."],
        ],
        widths=[1900, 3850, 3610],
    )
    writer.h2("Hiérarchie métier")
    writer.steps([
        "Le PAS fixe la stratégie et les objectifs de haut niveau.",
        "Le PAO rattache ces objectifs à une direction et à un exercice.",
        "Le PTA récupère les objectifs opérationnels transmis au service.",
        "Les actions sont créées dans le PTA, puis suivies dans le module Actions.",
        "Les validations transforment les données de suivi en données officielles.",
        "Les KPI, alertes et rapports consolident les résultats selon le périmètre autorisé.",
    ])

    writer.h1("3. Rôles, périmètres et responsabilités")
    writer.image("roles", "Illustration 2 - Vue synthétique des rôles et périmètres")
    writer.table(
        ["Profil", "Périmètre habituel", "Actions principales"],
        [
            ["Super Admin", "Global", "Paramétrage profond, modules, rôles, workflows, référentiels, maintenance, templates."],
            ["Admin / Admin fonctionnel", "Global", "Gestion métier, utilisateurs, référentiels, gouvernance et consultation globale."],
            ["DG", "Global", "Supervision globale, arbitrages critiques, lecture consolidée, audit et alertes."],
            ["Planification / SCIQ", "Global", "Structuration PAS/PAO/PTA, contrôle transversal, corrections exceptionnelles, reporting."],
            ["Direction", "Direction", "Création et suivi du PAO, supervision des PTA et actions de sa direction."],
            ["Service / Chef unité", "Service", "Création du PTA, création des actions, affectation des RMO, validation chef."],
            ["Agent / RMO", "Actions assignées", "Saisie de l'avancement, justificatifs, difficultés, commentaires et soumission."],
            ["Cabinet / collaborateur", "Lecture globale selon habilitation", "Consultation du pilotage, du reporting, des alertes et de l'audit."],
        ],
        widths=[2300, 2100, 4960],
    )
    writer.callout(
        "Attention",
        "Les menus visibles dépendent du rôle, du périmètre et de la configuration des modules. Deux utilisateurs peuvent donc voir des menus différents.",
        fill=WARNING_HEX,
    )

    writer.h1("4. Connexion, sécurité et profil")
    writer.h2("Se connecter")
    writer.steps([
        "Ouvrir l'adresse de l'application fournie par l'administrateur.",
        "Saisir l'adresse e-mail ou l'identifiant autorisé.",
        "Saisir le mot de passe.",
        "Cliquer sur Connexion.",
        "En cas d'oubli du mot de passe, utiliser le parcours de réinitialisation si celui-ci est activé.",
    ])
    writer.h2("Mettre à jour son profil")
    writer.bullets([
        "Accéder au menu Profil ou Paramètres personnels.",
        "Vérifier le nom complet, l'adresse e-mail, le rôle, la direction et le service de rattachement.",
        "Ajouter ou remplacer la photo de profil si le champ est disponible.",
        "Modifier le mot de passe depuis la section Sécurité.",
        "Révoquer les sessions actives si une connexion semble suspecte.",
    ])
    writer.callout(
        "Bonne pratique",
        "Ne partagez jamais votre compte. Les actions de validation, les commentaires, les exports et les changements sensibles sont journalisés dans l'audit.",
    )

    writer.h1("5. Interface et navigation")
    writer.p(
        "L'espace de travail regroupe les modules par famille : Menu, Planification, Exécution, Pilotage, Administration et Plateforme. Les libellés peuvent être personnalisés par le Super Admin, mais la logique reste identique."
    )
    writer.image("navigation", "Illustration 3 - Carte des menus principaux")
    writer.table(
        ["Famille", "Modules", "Utilisation"],
        [
            ["Menu", "Pilotage, Mes tâches, Notifications", "Accès rapide aux synthèses, tâches ouvertes et messages système."],
            ["Planification", "PAS, PAO, PTA, Imports Excel", "Création et structuration de la planification."],
            ["Exécution", "Actions, Financement des actions", "Suivi, contrôle, validation, financement et justificatifs."],
            ["Pilotage", "Reporting, Alertes", "Analyse consolidée, exports et surveillance des écarts."],
            ["Administration", "Référentiels, Délégations, Rétention, API Docs, Audit", "Gestion de la donnée de base, traçabilité et gouvernance."],
            ["Plateforme", "Super Administration", "Paramétrage avancé de l'application."],
        ],
        widths=[1900, 3000, 4460],
    )
    writer.h2("Utiliser les listes")
    writer.bullets([
        "Les cartes de synthèse en haut des pages donnent un accès rapide aux statuts importants.",
        "Les filtres permettent de réduire la liste par recherche, direction, service, année, statut ou responsable selon le module.",
        "Les boutons Modifier, Clôturer, Archiver ou Supprimer apparaissent uniquement si votre rôle le permet.",
        "Les actions sensibles déclenchent généralement une confirmation avant exécution.",
    ])

    writer.h1("6. Tableau de bord et Mes tâches")
    writer.p(
        "Le tableau de bord affiche une synthèse adaptée au rôle connecté : indicateurs de performance, alertes importantes, tâches ouvertes, score personnel, graphiques et tableaux de suivi."
    )
    writer.h2("Mes tâches")
    writer.bullets([
        "Le bloc Mes tâches signale les éléments qui attendent une action de votre part.",
        "Chaque tâche contient un titre, un contexte, une échéance et un lien pour traiter directement l'élément.",
        "Les retards de validation sont rattachés au valideur attendu, pas à l'agent qui a soumis à temps.",
        "Une tâche peut concerner une correction demandée, une validation chef, un financement, une alerte ou une demande de modification.",
    ])
    writer.h2("Lire les indicateurs")
    writer.bullets([
        "Progression déclarée : avancement calculé à partir des saisies de suivi.",
        "Progression théorique : progression attendue selon le calendrier.",
        "Performance officielle : valeur validée et utilisée dans les rapports.",
        "Alertes : écarts, retards, criticités ou anomalies détectés dans le périmètre.",
    ])

    writer.h1("7. Module PAS")
    writer.h2("Créer un PAS")
    writer.steps([
        "Ouvrir Planification > PAS.",
        "Cliquer sur Nouveau PAS.",
        "Renseigner le titre du PAS.",
        "Renseigner la période de début et la période de fin.",
        "Ajouter les axes stratégiques.",
        "Sous chaque axe, ajouter au moins un objectif stratégique avec sa date d'échéance.",
        "Cliquer sur Créer ou Mettre à jour.",
    ])
    writer.h2("Champs du PAS")
    writer.table(
        ["Champ", "Obligatoire", "Explication"],
        [
            ["Titre du PAS", "Oui", "Nom du plan stratégique."],
            ["Période début", "Oui", "Année de début du PAS."],
            ["Période fin", "Oui", "Année de fin du PAS."],
            ["Axe stratégique", "Oui", "Grand axe de la stratégie institutionnelle."],
            ["Objectif stratégique", "Oui", "Objectif rattaché à un axe."],
            ["Date d'échéance", "Oui", "Date limite associée à l'objectif stratégique."],
        ],
        widths=[2300, 1700, 5360],
    )
    writer.h2("Clôturer et archiver")
    writer.bullets([
        "Un PAS peut être clôturé quand les anomalies bloquantes ont été contrôlées.",
        "Le rapport d'anomalies peut signaler des PAO ouverts, PTA ouverts, actions en cours, validations en attente, retards ou KPI incomplets.",
        "L'archivage intervient après clôture et conserve la traçabilité.",
    ])

    writer.h1("8. Module PAO")
    writer.h2("Créer un PAO")
    writer.steps([
        "Ouvrir Planification > PAO.",
        "Cliquer sur Nouveau PAO.",
        "Sélectionner l'axe stratégique.",
        "Sélectionner l'objectif stratégique rattaché.",
        "Vérifier le PAS parent affiché en lecture seule.",
        "Sélectionner la direction concernée.",
        "Indiquer l'année de l'exercice.",
        "Ajouter un ou plusieurs objectifs opérationnels.",
        "Pour chaque objectif opérationnel, sélectionner le service concerné et l'échéance.",
        "Enregistrer le PAO.",
    ])
    writer.h2("Règles à respecter")
    writer.bullets([
        "Le PAO est directionnel : il est rattaché à une direction.",
        "Le service n'est pas porté par le PAO racine, mais par chaque objectif opérationnel.",
        "Une direction dispose en principe d'un PAO par exercice.",
        "Lorsque les champs obligatoires sont complets, la validation peut être automatique selon les règles configurées.",
        "Les objectifs opérationnels validés sont transmis aux chefs de service concernés.",
    ])
    writer.table(
        ["Champ PAO", "Usage"],
        [
            ["Axe stratégique", "Filtre l'objectif stratégique disponible."],
            ["Objectif stratégique", "Relie le PAO au PAS."],
            ["Direction", "Définit le périmètre directionnel du PAO."],
            ["Année", "Exercice de planification."],
            ["Objectif opérationnel", "Déclinaison concrète pour un service."],
            ["Service concerné", "Service destinataire de l'objectif opérationnel."],
            ["Échéance", "Date limite qui borne les actions du PTA."],
        ],
        widths=[2600, 6760],
    )

    writer.h1("9. Module PTA et création des actions")
    writer.p(
        "Le PTA est le point central de création des actions. Le module Actions sert ensuite au suivi, au contrôle, à la validation et à la consultation. Si un utilisateur cherche à créer une action directement depuis Actions, l'application le redirige vers le PTA."
    )
    writer.image("creation_action", "Illustration 4 - Les étapes de création d'une action dans le PTA")
    writer.h2("Créer un PTA")
    writer.steps([
        "Ouvrir Planification > PTA.",
        "Cliquer sur Nouveau PTA.",
        "Sélectionner l'objectif opérationnel transmis au service.",
        "Vérifier les informations affichées automatiquement : PAO d'origine, PAS lié, axe stratégique, objectif stratégique, direction, service et échéance.",
        "Créer les actions liées à l'objectif opérationnel.",
        "Enregistrer le PTA.",
    ])
    writer.h2("Créer une action dans le PTA")
    writer.steps([
        "Déplier le bloc Nouvelle action.",
        "Renseigner le libellé de l'action.",
        "Choisir le ou les RMO / agents assignés.",
        "Indiquer la date de début et la date de fin.",
        "Choisir le type d'action : quantitative, non quantitative ou composée.",
        "Renseigner la cible, l'unité et les seuils si l'action est quantitative.",
        "Ajouter les sous-actions si l'action est composée.",
        "Définir les ressources nécessaires, les risques et le financement si nécessaire.",
        "Cliquer sur Enregistrer pour sauvegarder l'action.",
    ])
    writer.table(
        ["Type d'action", "Quand l'utiliser", "Calcul"],
        [
            ["Quantitative", "Quand une cible chiffrée existe : nombre de dossiers, montants, formations, livrables quantifiés.", "Quantité réalisée / quantité cible."],
            ["Non quantitative", "Quand la réalisation se prouve surtout par une pièce ou un livrable.", "0 % sans preuve, 100 % provisoire avec preuve, officiel après validation."],
            ["Composée", "Quand l'action doit être découpée en sous-actions affectées à plusieurs RMO.", "Somme des performances de sous-actions pondérées."],
        ],
        widths=[2100, 4300, 2960],
    )
    writer.callout(
        "Point de vigilance",
        "Pour une action composée, la somme des poids des sous-actions doit être égale à 100 %. Cette règle garantit un calcul cohérent de la performance.",
        fill=WARNING_HEX,
    )

    writer.h1("10. Suivi d'une action")
    writer.image("workflow_action", "Illustration 5 - Workflow de suivi et validation d'une action")
    writer.h2("Ouvrir une action")
    writer.steps([
        "Ouvrir Exécution > Actions.",
        "Filtrer si nécessaire par statut, direction, service, responsable ou recherche.",
        "Cliquer sur Suivi ou ouvrir le détail de l'action.",
        "Consulter la fiche, la progression, les justificatifs, la discussion et le journal.",
    ])
    writer.h2("Enregistrer l'avancement")
    writer.bullets([
        "Pour une action quantitative, renseigner la quantité réalisée totale à ce jour.",
        "Pour une action non quantitative, déposer la pièce justificative attendue.",
        "Pour une sous-action, renseigner l'avancement ou la preuve dans le bloc de la sous-action.",
        "Ajouter un commentaire si nécessaire ou si le PTA l'a rendu obligatoire.",
        "Décrire les difficultés rencontrées si le champ est activé.",
        "Cliquer sur Enregistrer pour garder un brouillon sans déclencher la validation complète.",
    ])
    writer.h2("Soumettre au chef")
    writer.steps([
        "Vérifier que les champs requis par le PTA sont remplis.",
        "Déposer la pièce justificative si elle est obligatoire ou attendue.",
        "Compléter le commentaire si le commentaire est obligatoire.",
        "Renseigner les difficultés ou écrire Aucune difficulté rencontrée si demandé par la procédure interne.",
        "Cliquer sur Soumettre au chef.",
        "Attendre la validation ou la demande de correction.",
    ])
    writer.callout(
        "Différence essentielle",
        "Enregistrer calcule une performance provisoire. Soumettre déclenche le contrôle. Valider par le chef fige la performance officielle.",
    )

    writer.h1("11. Validation, corrections et demandes de modification")
    writer.h2("Validation chef")
    writer.bullets([
        "Le chef examine les éléments soumis par l'agent ou le RMO.",
        "Il vérifie la cohérence de la quantité, des justificatifs, des commentaires et des difficultés.",
        "S'il valide, l'action ou la sous-action devient officiellement prise en compte.",
        "S'il rejette, il doit renseigner un motif. L'élément revient en correction.",
    ])
    writer.h2("Corrections demandées")
    writer.steps([
        "Ouvrir Mes tâches ou le module Actions.",
        "Repérer l'action en correction demandée ou rejetée.",
        "Lire le motif de rejet dans la discussion ou le journal.",
        "Corriger la saisie, le justificatif, le commentaire ou la quantité.",
        "Enregistrer puis soumettre à nouveau.",
    ])
    writer.h2("Demande de modification")
    writer.bullets([
        "Après enregistrement définitif, certaines actions peuvent être figées en lecture seule.",
        "Le bouton Demande de modification permet de demander la réouverture.",
        "La demande suit le circuit contrôleur SCIQ/Planification puis décision DG selon la configuration.",
        "Le motif doit être clair : erreur de saisie, changement d'échéance, réaffectation, ajustement de cible ou correction de financement.",
    ])
    writer.h2("Financement")
    writer.bullets([
        "Si l'action nécessite un financement, le PTA doit indiquer le besoin, le montant, la nature et la pièce justificative.",
        "Le DAF peut valider et transmettre à la DG, demander un complément ou rejeter.",
        "La DG peut accorder ou refuser le financement.",
        "Les décisions, dates, pièces et commentaires restent visibles dans le détail de l'action.",
    ])

    writer.h1("12. Reporting, alertes et exports")
    writer.image("reporting_alertes", "Illustration 6 - Reporting, alertes et exports")
    writer.h2("Consulter un rapport")
    writer.steps([
        "Ouvrir Pilotage > Reporting.",
        "Choisir le type de rapport métier.",
        "Appliquer les filtres : exercice, trimestre, direction, service, statut, type d'action, responsable ou criticité.",
        "Analyser les résultats affichés.",
        "Exporter en Excel pour analyse détaillée ou en PDF pour diffusion.",
    ])
    writer.h2("Comprendre la base statistique")
    writer.bullets([
        "Le reporting respecte le périmètre autorisé de l'utilisateur connecté.",
        "Les chiffres officiels reposent sur les actions dont le circuit de validation requis est terminé.",
        "Les filtres modifient les tableaux et les exports.",
        "Les exports doivent être relus avant diffusion institutionnelle.",
    ])
    writer.h2("Gérer les alertes")
    writer.bullets([
        "Les alertes signalent les retards, criticités, anomalies et éléments à traiter.",
        "Une alerte peut être lue depuis Notifications > Alertes ou depuis le menu Alertes.",
        "Le bouton Ouvrir mène généralement vers l'action ou l'élément concerné.",
        "Les alertes non lues peuvent être marquées comme lues individuellement ou en masse selon les droits.",
    ])

    writer.h1("13. Référentiels, délégations, audit et gouvernance")
    writer.h2("Référentiels")
    writer.bullets([
        "Directions : créer, modifier ou désactiver les directions selon les droits.",
        "Services : rattacher chaque service à une direction.",
        "Utilisateurs : gérer le nom, l'e-mail, le rôle, le périmètre direction/service, le statut actif et la photo.",
        "Les suppressions sensibles peuvent suivre une demande de suppression plutôt qu'une suppression directe.",
    ])
    writer.h2("Délégations")
    writer.bullets([
        "Une délégation donne temporairement à un autre utilisateur la capacité d'intervenir dans un circuit.",
        "Elle doit préciser le délégant, le délégué, la période, le périmètre et le motif.",
        "Une délégation doit être annulée dès qu'elle n'est plus nécessaire.",
    ])
    writer.h2("Audit")
    writer.bullets([
        "Le journal d'audit conserve les actions sensibles : création, modification, suppression, validation, décisions et changements de configuration.",
        "Les filtres permettent de rechercher par module, action, utilisateur, entité, date ou texte.",
        "L'audit sert à contrôler la traçabilité et à comprendre l'historique d'un dossier.",
    ])
    writer.h2("Rétention et documentation API")
    writer.bullets([
        "La rétention concerne l'archivage et les règles de conservation des données.",
        "La documentation API expose les contrats techniques pour les intégrations autorisées.",
        "Ces modules sont réservés aux profils habilités.",
    ])

    writer.h1("14. Super Administration")
    writer.image("super_admin", "Illustration 7 - Zones de paramétrage de la Super Administration")
    writer.h2("Plateforme")
    writer.bullets([
        "Généraux : textes, logos, paramètres d'identité et formats.",
        "Apparence : palette, densité et options de lecture.",
        "Modules : visibilité, ordre et libellés des menus.",
        "Maintenance : actions techniques, caches et contrôles ponctuels.",
    ])
    writer.h2("Gouvernance")
    writer.bullets([
        "Rôles : matrice de permissions, registre des rôles et restauration de versions.",
        "Organisation : directions, services, comptes utilisateurs et import en masse.",
        "Unités DG : SCIQ, DGA, Cabinet, UCAS, chefs d'unité et membres.",
        "Dashboards : cartes et visibilité selon les profils.",
        "Diagnostic et audit : contrôle plateforme et actions sensibles.",
    ])
    writer.h2("Pilotage")
    writer.bullets([
        "Workflow : circuits Actions, PAS, PAO et PTA.",
        "Exercices : périodes, exercice actif et archivage automatique.",
        "Calcul : base statistique et règles de calcul des actions.",
        "Actions : paramètres métier de clôture, suspension et suivi.",
        "Référentiels dynamiques : libellés, unités, priorités et listes configurables.",
        "Documents : formats acceptés, rétention et droits liés aux justificatifs.",
        "Indicateur de performance : registre KPI et moteur de calcul.",
        "Notifications : événements, escalades et délais.",
    ])
    writer.h2("Avancé")
    writer.bullets([
        "Snapshots : sauvegarder, comparer et restaurer une configuration.",
        "Simulation : vérifier l'impact d'un changement avant application.",
        "Templates d'export : créer, prévisualiser, publier, archiver et affecter des modèles de rapports.",
    ])
    writer.callout(
        "Prudence",
        "Toute modification dans Super Administration peut changer l'expérience de plusieurs profils. Documentez le motif et utilisez les brouillons, snapshots ou simulations lorsqu'ils sont disponibles.",
        fill=WARNING_HEX,
    )

    writer.h1("15. Imports Excel")
    writer.h2("Importer un fichier")
    writer.steps([
        "Ouvrir Planification > Imports Excel.",
        "Télécharger le modèle Excel si nécessaire.",
        "Préparer une feuille avec une ligne par action planifiée.",
        "Cliquer sur Nouvel import.",
        "Choisir le fichier .xlsx ou .csv.",
        "Cliquer sur Vérifier le fichier.",
        "Analyser la prévisualisation : lignes valides, erreurs et avertissements.",
        "Corriger le fichier si nécessaire ou confirmer l'import.",
    ])
    writer.h2("Conseils de préparation")
    writer.bullets([
        "Respecter les intitulés et formats du modèle.",
        "Éviter les cellules fusionnées.",
        "Vérifier les codes directions, services, objectifs et utilisateurs avant import.",
        "Contrôler les dates et les montants.",
        "Lire le rapport d'erreurs si l'import échoue.",
    ])

    writer.h1("16. Bonnes pratiques par profil")
    writer.table(
        ["Profil", "Bonnes pratiques"],
        [
            ["Agent / RMO", "Mettre à jour l'avancement régulièrement, déposer les preuves, expliquer les difficultés, soumettre dès que les conditions sont remplies."],
            ["Chef de service", "Créer des actions claires, affecter les bons RMO, vérifier les cibles, valider ou rejeter rapidement avec un motif utile."],
            ["Direction", "Contrôler la cohérence du PAO, suivre les PTA de la direction, traiter les alertes de retard et arbitrer les priorités."],
            ["Planification / SCIQ", "Surveiller les anomalies, accompagner les corrections, maintenir la cohérence PAS-PAO-PTA et consolider le reporting."],
            ["DG", "Consulter les synthèses globales, arbitrer les cas critiques, suivre les alertes majeures et les demandes de modification sensibles."],
            ["Super Admin", "Utiliser les snapshots avant changements importants, tester les workflows, vérifier l'audit et limiter les droits sensibles."],
        ],
        widths=[2400, 6960],
    )

    writer.h1("17. Problèmes fréquents et solutions")
    writer.table(
        ["Situation", "Cause probable", "Solution"],
        [
            ["Je ne vois pas un module", "Votre rôle ou la configuration de navigation ne l'autorise pas.", "Demander à l'administrateur de vérifier vos droits et modules visibles."],
            ["Je ne peux pas créer une action depuis Actions", "Les actions se créent depuis le PTA.", "Aller dans Planification > PTA puis ajouter l'action dans le PTA."],
            ["Je ne peux pas soumettre", "Un champ obligatoire manque : preuve, commentaire, difficulté ou quantité.", "Lire les messages d'erreur, compléter les champs, puis soumettre à nouveau."],
            ["Mon action affiche 100 % mais n'est pas dans le reporting", "La performance est encore provisoire.", "Attendre ou demander la validation du chef."],
            ["Je dois modifier une action figée", "L'action est enregistrée et verrouillée.", "Utiliser Demande de modification et saisir un motif clair."],
            ["Un export ne correspond pas à mon attendu", "Les filtres ou le périmètre changent le résultat.", "Vérifier les filtres, l'exercice, le trimestre, la direction et le service."],
            ["Un justificatif est refusé", "Format non autorisé ou pièce incomplète.", "Utiliser un format accepté et déposer une pièce lisible."],
        ],
        widths=[2600, 3300, 3460],
    )

    writer.h1("18. Glossaire")
    writer.table(
        ["Terme", "Définition"],
        [
            ["ANBG", "Agence Nationale des Bourses du Gabon."],
            ["PAS", "Plan d'Actions Stratégique."],
            ["PAO", "Plan d'Actions Opérationnel."],
            ["PTA", "Plan de Travail Annuel."],
            ["RMO", "Responsable de mise en œuvre d'une action ou sous-action."],
            ["KPI", "Indicateur de performance."],
            ["Justificatif", "Document ou preuve déposée pour appuyer une réalisation."],
            ["Soumission", "Transmission d'une réalisation au chef pour validation."],
            ["Validation chef", "Décision qui officialise la performance."],
            ["Alerte", "Signal automatique ou manuel indiquant un risque, un retard ou une anomalie."],
            ["Snapshot", "Copie d'une configuration permettant comparaison ou restauration."],
            ["Périmètre", "Champ de visibilité et d'action autorisé pour un utilisateur."],
        ],
        widths=[2300, 7060],
    )

    writer.h1("19. Références internes utilisées")
    writer.bullets([
        "README.md : périmètre fonctionnel, comptes de démonstration et routes principales.",
        "GUIDE.md : organisation fonctionnelle et technique de l'application.",
        "docs/spec-workflow-canonique-pas-anbg.md : règles métier PAS, PAO, PTA, actions et validation.",
        "docs/WORKFLOW-SUIVI-V2.md : principes de suivi, performance provisoire/officielle et validation chef.",
        "Routes et vues Laravel : menus, formulaires, champs et écrans disponibles dans l'application.",
    ])


def save_markdown(writer: ManualWriter) -> None:
    MD_PATH.write_text("\n".join(writer.md).rstrip() + "\n", encoding="utf-8")


def main() -> None:
    DOCS.mkdir(exist_ok=True)
    assets = create_assets()
    doc = configure_document()
    writer = ManualWriter(doc, assets)
    add_cover(writer)
    build_manual(writer)
    doc.save(DOCX_PATH)
    save_markdown(writer)
    print(f"Created {DOCX_PATH}")
    print(f"Created {MD_PATH}")


if __name__ == "__main__":
    main()
