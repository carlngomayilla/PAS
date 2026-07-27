from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DOCUMENTS = [
    (
        DOCS / "cahier-specifications-fonctionnelles-application.md",
        DOCS / "cahier-specifications-fonctionnelles-application.docx",
        "Cahier de spécifications fonctionnelles",
        "Référentiel fonctionnel de l'application ANBG e-Pilotage",
    ),
    (
        DOCS / "analyse-fonctionnelle-workflows.md",
        DOCS / "analyse-fonctionnelle-workflows.docx",
        "Analyse fonctionnelle des workflows métier",
        "Circuits PAS, PAO, PTA, actions, gouvernance et reporting",
    ),
]

NAVY = RGBColor(11, 55, 86)
BLUE = RGBColor(28, 117, 180)
MUTED = RGBColor(91, 105, 120)
LIGHT_BLUE = "EAF4FB"
LIGHT_GRAY = "F3F6F8"


def clean_inline(text: str) -> str:
    text = re.sub(r"!\[[^\]]*]\([^)]+\)", "", text)
    text = re.sub(r"\[([^\]]+)]\([^)]+\)", r"\1", text)
    text = text.replace("**", "").replace("__", "").replace("`", "")
    return text.strip()


def shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, value: int = 110) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name in ("top", "start", "bottom", "end"):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    properties = table._tbl.tblPr
    width = properties.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(sum(widths)))
    width.set(qn("w:type"), "dxa")

    indent = properties.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for column_width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(column_width))
        grid.append(column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            properties = cell._tc.get_or_add_tcPr()
            cell_width = properties.find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                properties.append(cell_width)
            cell_width.set(qn("w:w"), str(widths[index]))
            cell_width.set(qn("w:type"), "dxa")


def finalize_accessibility(document: Document) -> None:
    for table in document.tables:
        properties = table.rows[0]._tr.get_or_add_trPr()
        header = properties.find(qn("w:tblHeader"))
        if header is None:
            header = OxmlElement("w:tblHeader")
            properties.append(header)
        header.set(qn("w:val"), "true")


def configure_document(title: str, subtitle: str) -> Document:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in [
        ("Heading 1", 16, NAVY, 16, 7),
        ("Heading 2", 13, BLUE, 12, 5),
        ("Heading 3", 11, NAVY, 9, 4),
    ]:
        style = document.styles[style_name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "ANBG e-Pilotage - Documentation fonctionnelle"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8.5)
        run.font.color.rgb = MUTED

    footer = section.footer.paragraphs[0]
    footer.text = "Version consolidée du 27 juillet 2026 - Usage interne"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8.5)
        run.font.color.rgb = MUTED

    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_before = Pt(100)
    title_run = title_paragraph.add_run(title)
    title_run.font.name = "Aptos Display"
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = NAVY

    subtitle_paragraph = document.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_paragraph.add_run(subtitle)
    subtitle_run.font.name = "Aptos"
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = MUTED

    document.add_paragraph()
    metadata = document.add_table(rows=3, cols=2)
    metadata.style = "Table Grid"
    values = [
        ("Application", "ANBG e-Pilotage PAS / PAO / PTA / Actions"),
        ("Référence", "Code Laravel, routes, services, modèles, tests et écrans actuels"),
        ("Mise à jour", "27 juillet 2026"),
    ]
    for row, (label, value) in zip(metadata.rows, values):
        shade_cell(row.cells[0], LIGHT_BLUE)
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        label_run = row.cells[0].paragraphs[0].add_run(label)
        label_run.bold = True
        label_run.font.color.rgb = NAVY
        row.cells[1].paragraphs[0].add_run(value)
    set_table_geometry(metadata, [2500, 6860])
    document.add_page_break()
    return document


def add_markdown_table(document: Document, lines: list[str]) -> None:
    rows = [[clean_inline(value) for value in line.strip().strip("|").split("|")] for line in lines]
    if len(rows) < 2:
        return
    data_rows = [rows[0]] + rows[2:]
    column_count = len(data_rows[0])
    table = document.add_table(rows=1, cols=column_count)
    table.style = "Table Grid"
    for index, value in enumerate(data_rows[0]):
        cell = table.rows[0].cells[index]
        shade_cell(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(value)
        run.bold = True
        run.font.color.rgb = NAVY
    for values in data_rows[1:]:
        cells = table.add_row().cells
        for index in range(column_count):
            value = values[index] if index < len(values) else ""
            set_cell_margins(cells[index])
            cells[index].paragraphs[0].add_run(value.replace("<br>", "\n"))
    base = 9360 // column_count
    widths = [base] * column_count
    widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths)
    document.add_paragraph()


def convert_markdown(source: Path, output: Path, title: str, subtitle: str) -> None:
    document = configure_document(title, subtitle)
    lines = source.read_text(encoding="utf-8").splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []

    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()

        if stripped.startswith("```"):
            if in_code:
                table = document.add_table(rows=1, cols=1)
                table.style = "Table Grid"
                shade_cell(table.cell(0, 0), LIGHT_GRAY)
                set_cell_margins(table.cell(0, 0), 140)
                paragraph = table.cell(0, 0).paragraphs[0]
                run = paragraph.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(8.5)
                set_table_geometry(table, [9360])
                document.add_paragraph()
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue

        if in_code:
            code_lines.append(raw)
            index += 1
            continue

        if stripped.startswith("|") and index + 1 < len(lines) and re.match(r"^\s*\|[\s:|-]+\|\s*$", lines[index + 1]):
            table_lines = [raw, lines[index + 1]]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(document, table_lines)
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            text = clean_inline(heading.group(2))
            if text not in {title, "Cahier de specifications fonctionnelles", "Analyse fonctionnelle des workflows métier"}:
                document.add_heading(text, level=len(heading.group(1)))
            index += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            text = clean_inline(re.sub(r"^\d+\.\s+", "", stripped))
            document.add_paragraph(text, style="List Number")
        elif stripped.startswith("- "):
            document.add_paragraph(clean_inline(stripped[2:]), style="List Bullet")
        elif stripped.startswith(">"):
            table = document.add_table(rows=1, cols=1)
            shade_cell(table.cell(0, 0), LIGHT_BLUE)
            set_cell_margins(table.cell(0, 0), 140)
            table.cell(0, 0).paragraphs[0].add_run(clean_inline(stripped.lstrip("> ")))
            set_table_geometry(table, [9360])
            document.add_paragraph()
        elif stripped in {"---", ""}:
            pass
        else:
            paragraph = document.add_paragraph()
            paragraph.add_run(clean_inline(stripped))
        index += 1

    finalize_accessibility(document)
    document.save(output)
    print(f"Created {output}")


def main() -> None:
    for source, output, title, subtitle in DOCUMENTS:
        convert_markdown(source, output, title, subtitle)


if __name__ == "__main__":
    main()
